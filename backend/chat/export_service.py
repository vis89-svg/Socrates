import io
import re
from html import escape

from docx import Document
from docx.shared import Pt


def _iter_blocks(markdown):
    """Split markdown into logical blocks: heading, table, list, hr, or paragraph."""
    lines = markdown.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            yield ('table', table_lines)
            continue
        if stripped.startswith('```'):
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i])
                i += 1
            i += 1
            yield ('code', '\n'.join(code))
            continue
        if re.match(r'^#{1,6}\s', stripped):
            yield ('heading', stripped)
            i += 1
            continue
        if re.match(r'^[-*+]\s+', stripped):
            list_lines = []
            while i < n and re.match(r'^[-*+]\s+', lines[i].strip()):
                list_lines.append(lines[i].strip()[2:].strip())
                i += 1
            yield ('list', list_lines)
            continue
        if re.match(r'^\d+[.)]\s+', stripped):
            list_lines = []
            while i < n and re.match(r'^\d+[.)]\s+', lines[i].strip()):
                list_lines.append(lines[i].strip())
                i += 1
            yield ('list', list_lines)
            continue
        if re.match(r'^-{3,}$', stripped):
            yield ('hr', None)
            i += 1
            continue
        para = [line]
        i += 1
        while i < n and lines[i].strip() and not (
            re.match(r'^#{1,6}\s', lines[i].strip())
            or re.match(r'^[-*+]\s+', lines[i].strip())
            or re.match(r'^\d+[.)]\s+', lines[i].strip())
            or lines[i].strip().startswith('|')
            or lines[i].strip().startswith('```')
            or re.match(r'^-{3,}$', lines[i].strip())
        ):
            para.append(lines[i])
            i += 1
        yield ('paragraph', ' '.join(p.strip() for p in para))


def _parse_table(lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        if all(re.match(r'^:?-{2,}:?$', c) for c in cells if c):
            continue
        rows.append(cells)
    return rows


def _inline_md(text):
    """Strip inline markdown (bold, italics, links, code) for plain-text contexts."""
    text = re.sub(r'!\[([^\]]*)\]\([^)]*\)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


class ExportService:
    @staticmethod
    def to_docx(markdown):
        doc = Document()
        doc.add_heading('Socrates Chat Report', level=0)
        for kind, content in _iter_blocks(markdown):
            if kind == 'heading':
                level = min(len(content) - len(content.lstrip('#')), 6)
                text = _inline_md(content.lstrip('# ').strip())
                doc.add_heading(text, level=level)
            elif kind == 'paragraph':
                doc.add_paragraph(_inline_md(content))
            elif kind == 'list':
                for item in content:
                    doc.add_paragraph(_inline_md(item), style='List Bullet')
            elif kind == 'code':
                for line in content.splitlines():
                    doc.add_paragraph(line, style='No Spacing')
            elif kind == 'hr':
                doc.add_paragraph('─' * 40)
            elif kind == 'table':
                rows = _parse_table(content)
                if not rows:
                    continue
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = _inline_md(row[c_idx]) if c_idx < len(row) else ''
                        if r_idx == 0:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    @staticmethod
    def to_pdf(markdown):
        import fitz

        html_parts = ['<html><head><style>'
                      'body{font-family:helvetica;font-size:10pt;color:#1c1e21;line-height:1.4;}'
                      'h1{font-size:17pt;} h2{font-size:14pt;} h3{font-size:12pt;}'
                      'h4,h5,h6{font-size:11pt;}'
                      'table{border-collapse:collapse;margin:6pt 0;width:100%;}'
                      'td,th{border:0.5pt solid #999;padding:3pt 5pt;font-size:9pt;}'
                      'th{background:#eef1f5;font-weight:bold;}'
                      'code{font-family:courier;font-size:9pt;background:#f4f4f4;}'
                      'pre{background:#f4f4f4;padding:6pt;font-family:courier;font-size:9pt;}'
                      'hr{border:none;border-top:0.5pt solid #ccc;margin:8pt 0;}'
                      '</style></head><body>']
        for kind, content in _iter_blocks(markdown):
            if kind == 'heading':
                level = min(len(content) - len(content.lstrip('#')), 6)
                text = escape(content.lstrip('# ').strip())
                html_parts.append(f'<h{level}>{text}</h{level}>')
            elif kind == 'paragraph':
                html_parts.append(f'<p>{ExportService._inline_html(content)}</p>')
            elif kind == 'list':
                items = ''.join(f'<li>{ExportService._inline_html(item)}</li>' for item in content)
                ordered = bool(re.match(r'^\d', content[0])) if content else False
                tag = 'ol' if ordered else 'ul'
                html_parts.append(f'<{tag}>{items}</{tag}>')
            elif kind == 'code':
                html_parts.append(f'<pre>{escape(content)}</pre>')
            elif kind == 'hr':
                html_parts.append('<hr/>')
            elif kind == 'table':
                rows = _parse_table(content)
                if not rows:
                    continue
                html_rows = []
                for r_idx, row in enumerate(rows):
                    tag = 'th' if r_idx == 0 else 'td'
                    cells = ''.join(f'<{tag}>{escape(c)}</{tag}>' for c in row)
                    html_rows.append(f'<tr>{cells}</tr>')
                html_parts.append('<table>' + ''.join(html_rows) + '</table>')
        html_parts.append('</body></html>')
        html = ''.join(html_parts)

        import fitz

        def rectfn(rect_num, filled):
            return fitz.Rect(0, 0, 595, 842), fitz.Rect(40, 40, 555, 802), None

        doc = fitz.Story(html=html).write_with_links(rectfn)
        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        buf.seek(0)
        return buf

    @staticmethod
    def _inline_html(text):
        text = escape(text)
        text = re.sub(r'!\[([^\]]*)\]\(([^)]*)\)', r'<img src="\2" alt="\1"/>', text)
        text = re.sub(r'\[([^\]]+)\]\(([^)]*)\)', r'<a href="\2">\1</a>', text)
        text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', text)
        text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
        text = re.sub(r'\n', '<br/>', text)
        return text
